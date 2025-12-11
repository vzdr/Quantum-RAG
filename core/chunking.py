"Document loading and chunking module."
import re
from pathlib import Path
from typing import List
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from .data_models import Document, Chunk

def load_document(file_path: str) -> Document:
    """Loads a document from a file path."""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    content = ""
    if ext == ".pdf":
        reader = PdfReader(file_path)
        content = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        content = "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    elif ext == ".txt":
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}")
        
    return Document(content=content, source=path.name, file_type=ext)

def chunk_document(document: Document, chunk_size: int = 500, overlap: int = 50) -> List[Chunk]:
    """Chunks a document into smaller pieces."""
    sentences = re.split(r'(?<=[.!?])\s+', document.content)
    sentences = [s.strip() for s in sentences if s.strip()]
    
    chunks = []
    current_chunk_sentences = []
    current_length = 0
    
    for sentence in sentences:
        if current_length + len(sentence) > chunk_size and current_chunk_sentences:
            text = " ".join(current_chunk_sentences)
            chunks.append(Chunk(id=f"{document.source}_{len(chunks)}", text=text, source=document.source))
            
            # Handle overlap
            overlap_sentences = []
            overlap_len = 0
            for s in reversed(current_chunk_sentences):
                if overlap_len + len(s) <= overlap:
                    overlap_sentences.insert(0, s)
                    overlap_len += len(s) + 1
                else:
                    break
            current_chunk_sentences = overlap_sentences
            current_length = overlap_len
        
        current_chunk_sentences.append(sentence)
        current_length += len(sentence)
        
    if current_chunk_sentences:
        text = " ".join(current_chunk_sentences)
        chunks.append(Chunk(id=f"{document.source}_{len(chunks)}", text=text, source=document.source))
        
    return chunks
