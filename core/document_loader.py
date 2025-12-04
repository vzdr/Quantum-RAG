"""
Document Loader Module - Phase A, Step 1

Handles loading documents from various file formats:
- PDF (.pdf)
- Plain text (.txt)
- Word documents (.docx)
"""
import os
import tempfile
from pathlib import Path
from dataclasses import dataclass, field
from typing import Dict, Any, List, Optional

import chardet
from PyPDF2 import PdfReader
from docx import Document as DocxDocument


@dataclass
class Document:
    """Represents a loaded document."""
    content: str
    source: str
    file_type: str
    metadata: Dict[str, Any] = field(default_factory=dict)

    def __len__(self) -> int:
        """Return the length of the document content."""
        return len(self.content)

    def preview(self, max_chars: int = 500) -> str:
        """Return a preview of the document content."""
        if len(self.content) <= max_chars:
            return self.content
        return self.content[:max_chars] + "..."


class DocumentLoader:
    """
    Handles loading documents from various file formats.

    Supported formats:
    - PDF (.pdf)
    - Plain text (.txt)
    - Word documents (.docx)
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.txt', '.docx'}

    @staticmethod
    def load_pdf(file_path: str) -> Document:
        """
        Load text from a PDF file.

        Args:
            file_path: Path to the PDF file

        Returns:
            Document object with extracted text
        """
        reader = PdfReader(file_path)
        text_parts = []

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        text = "\n".join(text_parts)

        return Document(
            content=text,
            source=Path(file_path).name,
            file_type="pdf",
            metadata={
                "page_count": len(reader.pages),
                "file_path": str(file_path),
                "char_count": len(text),
            }
        )

    @staticmethod
    def load_txt(file_path: str) -> Document:
        """
        Load text from a TXT file.

        Args:
            file_path: Path to the text file

        Returns:
            Document object with file content
        """
        # Detect encoding
        with open(file_path, 'rb') as f:
            raw = f.read()
            detected = chardet.detect(raw)
            encoding = detected.get('encoding', 'utf-8') or 'utf-8'

        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            text = f.read()

        return Document(
            content=text,
            source=Path(file_path).name,
            file_type="txt",
            metadata={
                "encoding": encoding,
                "file_path": str(file_path),
                "char_count": len(text),
                "line_count": text.count('\n') + 1,
            }
        )

    @staticmethod
    def load_docx(file_path: str) -> Document:
        """
        Load text from a DOCX file.

        Args:
            file_path: Path to the Word document

        Returns:
            Document object with extracted text
        """
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n\n".join(paragraphs)

        return Document(
            content=text,
            source=Path(file_path).name,
            file_type="docx",
            metadata={
                "paragraph_count": len(paragraphs),
                "file_path": str(file_path),
                "char_count": len(text),
            }
        )

    @classmethod
    def load_from_bytes(cls, content: bytes, filename: str) -> Document:
        """
        Load document from bytes (for widget upload).

        Args:
            content: File content as bytes
            filename: Original filename

        Returns:
            Document object
        """
        suffix = Path(filename).suffix.lower()

        if suffix not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {suffix}. Supported: {cls.SUPPORTED_EXTENSIONS}")

        # Write to temporary file and load
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        try:
            return cls.load(tmp_path)
        finally:
            os.unlink(tmp_path)

    @classmethod
    def load(cls, file_path: str) -> Document:
        """
        Load a document based on its file extension.

        Args:
            file_path: Path to the document

        Returns:
            Document object

        Raises:
            ValueError: If file type is not supported
            FileNotFoundError: If file does not exist
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = path.suffix.lower()

        loaders = {
            '.pdf': cls.load_pdf,
            '.txt': cls.load_txt,
            '.docx': cls.load_docx,
        }

        if ext not in loaders:
            raise ValueError(f"Unsupported file type: {ext}. Supported: {list(loaders.keys())}")

        return loaders[ext](file_path)

    @classmethod
    def load_directory(cls, directory_path: str, recursive: bool = False) -> List[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory_path: Path to the directory
            recursive: Whether to search subdirectories

        Returns:
            List of Document objects
        """
        documents = []
        path = Path(directory_path)

        if not path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")

        pattern = '**/*' if recursive else '*'

        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in cls.SUPPORTED_EXTENSIONS:
                try:
                    doc = cls.load(str(file_path))
                    documents.append(doc)
                except Exception as e:
                    print(f"Warning: Failed to load {file_path}: {e}")

        return documents

    @staticmethod
    def get_file_info(file_path: str) -> Dict[str, Any]:
        """
        Get file information without loading content.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with file information
        """
        path = Path(file_path)
        stat = path.stat()

        return {
            "name": path.name,
            "extension": path.suffix.lower(),
            "size_bytes": stat.st_size,
            "size_kb": round(stat.st_size / 1024, 2),
            "modified": stat.st_mtime,
        }
